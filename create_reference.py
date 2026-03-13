"""
Generate a reference.docx template for pandoc with professional styling.
Then convert all markdown files in docs/ to styled .docx files.
"""
import subprocess
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCS_DIR = "docs"
OUTPUT_DIR = "docs_docx"
REFERENCE_FILE = "reference.docx"

FONT_NAME = "Arial"
HEADING_COLOR = RGBColor(0x1A, 0x3C, 0x6E)  # Dark navy blue
ACCENT_COLOR = RGBColor(0x2E, 0x74, 0xB5)   # Medium blue


def set_font(run, name=FONT_NAME, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if italic is not None:
        run.font.italic = italic


def create_reference():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.line_spacing = 1.15
    rpr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
    rpr.append(rFonts)

    heading_configs = [
        ('Heading 1', 22, True, HEADING_COLOR, 18, 12),
        ('Heading 2', 16, True, HEADING_COLOR, 14, 8),
        ('Heading 3', 13, True, ACCENT_COLOR, 10, 6),
        ('Heading 4', 11.5, True, ACCENT_COLOR, 8, 4),
        ('Heading 5', 11, True, RGBColor(0x44, 0x44, 0x44), 6, 4),
    ]

    for style_name, size, bold, color, space_before, space_after in heading_configs:
        s = doc.styles[style_name]
        s.font.name = FONT_NAME
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(space_before)
        s.paragraph_format.space_after = Pt(space_after)
        s.paragraph_format.keep_with_next = True
        rpr = s.element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)

    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = FONT_NAME
    title_style.font.size = Pt(26)
    title_style.font.bold = True
    title_style.font.color.rgb = HEADING_COLOR
    title_style.paragraph_format.space_after = Pt(16)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    sub_style = doc.styles['Subtitle']
    sub_style.font.name = FONT_NAME
    sub_style.font.size = Pt(14)
    sub_style.font.color.rgb = ACCENT_COLOR
    sub_style.font.italic = True
    sub_style.paragraph_format.space_after = Pt(12)

    # List styles
    for list_style_name in ['List Bullet', 'List Number', 'List Paragraph']:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = FONT_NAME
            ls.font.size = Pt(11)
        except KeyError:
            pass

    # Table style  
    try:
        ts = doc.styles['Table Grid']
        ts.font.name = FONT_NAME
        ts.font.size = Pt(10)
    except KeyError:
        pass

    # Code style (for inline code)
    try:
        from docx.enum.style import WD_STYLE_TYPE
        code_char = doc.styles.add_style('Source Code', WD_STYLE_TYPE.CHARACTER)
        code_char.font.name = 'Courier New'
        code_char.font.size = Pt(10)
        code_char.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    except Exception:
        pass

    # Add placeholder content so pandoc picks up styles
    p = doc.add_paragraph("", style='Title')
    p = doc.add_paragraph("", style='Subtitle')
    for i in range(1, 6):
        doc.add_heading("", level=i)
    doc.add_paragraph("")

    doc.save(REFERENCE_FILE)
    print(f"Reference template created: {REFERENCE_FILE}")


def convert_all():
    md_files = sorted([f for f in os.listdir(DOCS_DIR) if f.endswith('.md')])
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    for md_file in md_files:
        input_path = os.path.join(DOCS_DIR, md_file)
        output_name = md_file.replace('.md', '.docx')
        output_path = os.path.join(OUTPUT_DIR, output_name)

        cmd = [
            'pandoc',
            input_path,
            '-o', output_path,
            '--reference-doc', REFERENCE_FILE,
            '--from', 'markdown+pipe_tables+grid_tables+multiline_tables+raw_html+emoji+strikeout',
            '--to', 'docx',
            '--toc',
            '--toc-depth=3',
            '--highlight-style', 'tango',
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            print(f"OK: {output_name}")
        else:
            print(f"ERROR: {md_file}: {result.stderr}")

    print(f"\nAll files saved to {OUTPUT_DIR}/")


def postprocess_docx():
    """Apply additional styling to tables in generated docx files."""
    for f in sorted(os.listdir(OUTPUT_DIR)):
        if not f.endswith('.docx'):
            continue
        path = os.path.join(OUTPUT_DIR, f)
        try:
            doc = Document(path)
            changed = False

            for table in doc.tables:
                changed = True
                table.alignment = WD_TABLE_ALIGNMENT.LEFT

                tbl = table._tbl
                tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

                borders_xml = f'''
                <w:tblBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>
                    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B0B0B0"/>
                </w:tblBorders>'''
                tbl_pr.append(parse_xml(borders_xml))

                if len(table.rows) > 0:
                    header_row = table.rows[0]
                    for cell in header_row.cells:
                        tc = cell._tc
                        tc_pr = tc.get_or_add_tcPr()
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="1A3C6E" w:val="clear"/>'
                        )
                        tc_pr.append(shading)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.size = Pt(10)
                                run.font.name = FONT_NAME

                for i, row in enumerate(table.rows):
                    if i == 0:
                        continue
                    bg = "F5F7FA" if i % 2 == 1 else "FFFFFF"
                    for cell in row.cells:
                        tc = cell._tc
                        tc_pr = tc.get_or_add_tcPr()
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'
                        )
                        tc_pr.append(shading)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                                run.font.name = FONT_NAME

                # Bold rows that contain "Итого"
                for row in table.rows:
                    row_text = " ".join(c.text for c in row.cells)
                    if "Итого" in row_text or "итого" in row_text or "ИТОГО" in row_text:
                        for cell in row.cells:
                            tc = cell._tc
                            tc_pr = tc.get_or_add_tcPr()
                            shading = parse_xml(
                                f'<w:shd {nsdecls("w")} w:fill="E8EDF5" w:val="clear"/>'
                            )
                            tc_pr.append(shading)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True

            if changed:
                doc.save(path)
                print(f"Styled tables: {f}")
            else:
                print(f"No tables: {f}")
        except Exception as e:
            print(f"Post-process error {f}: {e}")


if __name__ == '__main__':
    create_reference()
    convert_all()
    postprocess_docx()
